#!/usr/bin/env python3
"""
Обработка всех книг по продажам - FB2, PDF, DOC
"""

from pathlib import Path
import json
import re
import xml.etree.ElementTree as ET
import csv
import pandas as pd
from datetime import datetime

class SalesBookProcessor:
    """Процессор книг по продажам - извлекает техники из разных форматов"""
    
    def __init__(self):
        self.books_dir = Path("books")
        self.data_dir = Path("data")
        self.data_dir.mkdir(exist_ok=True)
        
        # Паттерны для поиска техник продаж
        self.technique_patterns = {
            'выявление_потребностей': [
                r'[А-ЯЁ][^.!?]*(?:что|как|почему|когда|где|какой|какая|какие)[^.!?]*\?',
                r'[А-ЯЁ][^.!?]*(?:расскажите|объясните|опишите|поделитесь)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:важно|интересует|волнует|беспокоит)[^.!?]*\?',
                r'[А-ЯЁ][^.!?]*(?:потребност|нужд|требован)[^.!?]*[.!?]'
            ],
            'презентация_объектов': [
                r'[А-ЯЁ][^.!?]*(?:это означает что|благодаря этому|в результате)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:представьте себе|посмотрите|обратите внимание)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:выгода|преимущество|польза)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:квартира|дом|объект)[^.!?]*(?:характеристик|особенност)[^.!?]*[.!?]'
            ],
            'работа_с_возражениями': [
                r'[А-ЯЁ][^.!?]*(?:понимаю|согласен)[^.!?]*(?:но|однако|тем не менее)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:возражение|сомнение)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:дорого)[^.!?]*[.!?]'
            ],
            'установление_доверия': [
                r'[А-ЯЁ][^.!?]*(?:опыт показывает|практика показывает)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:другие клиенты|наши клиенты)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:гарантирую|обещаю|ручаюсь)[^.!?]*[.!?]',
                r'[А-ЯЁ][^.!?]*(?:доверие|надежность)[^.!?]*[.!?]'
            ],
            'закрытие_сделки': [
                r'[А-ЯЁ][^.!?]*(?:готовы|согласны)[^.!?]*(?:подписать|оформить|купить)[^.!?]*\?',
                r'[А-ЯЁ][^.!?]*(?:когда удобно|когда вам подходит)[^.!?]*\?',
                r'[А-ЯЁ][^.!?]*(?:выбираете|принимаете решение)[^.!?]*\?'
            ]
        }
        
        # Паттерны для диалогов
        self.dialog_patterns = [
            r'(?:Менеджер|Продавец|Риелтор|Агент):\s*([^.!?]+[.!?])',
            r'(?:Клиент|Покупатель|Заказчик):\s*([^.!?]+[.!?])',
            r'—\s*([А-ЯЁ][^.!?]+[.!?])'
        ]
    
    def process_all_books(self):
        """Обрабатывает все книги в папке books (FB2, PDF, DOC) и извлекает техники продаж"""
        all_techniques = []
        
        # Находим все поддерживаемые файлы
        supported_extensions = ["*.fb2", "*.pdf", "*.doc", "*.docx"]
        all_files = []
        
        for ext in supported_extensions:
            all_files.extend(list(self.books_dir.glob(ext)))
        
        if not all_files:
            print("❌ Книги не найдены в папке books/")
            return []
        
        print(f"📚 Найдено файлов: {len(all_files)}")
        print(f"   FB2: {len(list(self.books_dir.glob('*.fb2')))}")
        print(f"   PDF: {len(list(self.books_dir.glob('*.pdf')))}")
        print(f"   DOC: {len(list(self.books_dir.glob('*.doc*')))}")
        
        for book_file in all_files:
            print(f"📖 Обрабатываю: {book_file.name}")
            
            # Извлекаем текст в зависимости от формата
            text = self.extract_text_from_file(book_file)
            
            if not text:
                print(f"   ⚠️ Не удалось извлечь текст")
                continue
            
            # Ищем техники
            book_techniques = self.extract_techniques(text, book_file.name)
            all_techniques.extend(book_techniques)
            
            print(f"   ✅ Найдено техник: {len(book_techniques)}")
        
        # Добавляем демо техники для базовой теории (только один раз)
        demo_techniques = self.get_demo_techniques()
        all_techniques.extend(demo_techniques)
        print(f"📝 Добавлено базовых техник: {len(demo_techniques)}")
        
        return all_techniques
    
    def extract_text_from_file(self, file_path):
        """Извлекает текст из файла в зависимости от его формата"""
        file_extension = file_path.suffix.lower()
        
        try:
            if file_extension == '.fb2':
                return self.extract_fb2_text(file_path)
            elif file_extension == '.pdf':
                return self.extract_pdf_text(file_path)
            elif file_extension in ['.doc', '.docx']:
                return self.extract_doc_text(file_path)
            else:
                print(f"   ⚠️ Неподдерживаемый формат: {file_extension}")
                return ""
        except Exception as e:
            print(f"   ❌ Ошибка обработки файла: {e}")
            return ""
    
    def extract_fb2_text(self, fb2_path):
        """Извлечение текста из FB2 файла"""
        methods = [
            ('xml_parse', self._xml_extract),
            ('regex_clean', self._regex_extract),
            ('binary_force', self._binary_extract)
        ]
        
        for method_name, method_func in methods:
            try:
                text = method_func(fb2_path)
                if text and len(text) > 1000:
                    return self.clean_text(text)
            except Exception:
                continue
        return ""

    def _xml_extract(self, fb2_path):
        """XML извлечение"""
        tree = ET.parse(fb2_path)
        root = tree.getroot()
        
        paragraphs = []
        for elem in root.iter():
            if elem.tag and elem.tag.endswith('p') and elem.text:
                text = elem.text.strip()
                if len(text) > 10:
                    paragraphs.append(text)
        
        return '\n\n'.join(paragraphs)

    def _regex_extract(self, fb2_path):
        """Regex извлечение"""
        for encoding in ['utf-8', 'cp1251', 'koi8-r']:
            try:
                with open(fb2_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                clean_text = re.sub(r'<[^>]*>', '\n', content)
                clean_text = re.sub(r'\n\s*\n', '\n\n', clean_text)
                
                if len(clean_text) > 2000:
                    return clean_text
            except:
                continue
        return ""

    def _binary_extract(self, fb2_path):
        """Принудительное извлечение"""
        with open(fb2_path, 'rb') as f:
            raw_data = f.read()
        
        for encoding in ['utf-8', 'cp1251']:
            try:
                content = raw_data.decode(encoding, errors='ignore')
                clean_text = re.sub(r'<[^>]*>', '\n', content)
                if len(clean_text) > 2000:
                    return clean_text
            except:
                continue
        return ""
    
    def extract_pdf_text(self, pdf_path):
        """Извлечение текста из PDF файла"""
        text = ""
        
        # Пробуем разные библиотеки для PDF
        try:
            # Сначала пробуем PyPDF2
            import PyPDF2
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            if len(text.strip()) > 100:
                return self.clean_text(text)
        except Exception:
            pass
        
        try:
            # Если PyPDF2 не сработал, пробуем pdfplumber
            import pdfplumber
            with pdfplumber.open(pdf_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if len(text.strip()) > 100:
                return self.clean_text(text)
        except Exception:
            pass
        
        # Если ничего не сработало, возвращаем пустую строку
        return ""
    
    def extract_doc_text(self, doc_path):
        """Извлечение текста из DOC/DOCX файла"""
        text = ""
        
        # Для DOCX файлов
        if doc_path.suffix.lower() == '.docx':
            try:
                from docx import Document
                doc = Document(doc_path)
                
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                return self.clean_text(text)
            except Exception as e:
                print(f"   ⚠️ Ошибка чтения DOCX: {e}")
        
        # Для старых DOC файлов пробуем разные методы
        else:
            try:
                # Пробуем через python-docx (иногда работает с .doc)
                from docx import Document
                doc = Document(doc_path)
                
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                
                if len(text.strip()) > 100:
                    return self.clean_text(text)
            except Exception:
                pass
            
            try:
                # Принудительное чтение как текст
                with open(doc_path, 'rb') as f:
                    raw_data = f.read()
                
                # Пробуем разные кодировки
                for encoding in ['utf-8', 'cp1251', 'latin1']:
                    try:
                        text = raw_data.decode(encoding, errors='ignore')
                        # Очищаем от мусора
                        clean_text = re.sub(r'[^\w\s.,!?;:()\-«»""„"„'']', ' ', text)
                        clean_text = re.sub(r'\s+', ' ', clean_text)
                        
                        if len(clean_text.strip()) > 1000:
                            return clean_text
                    except:
                        continue
            except Exception:
                pass
        
        return ""
    
    def extract_techniques(self, text, book_name):
        """Извлекает техники продаж из текста"""
        techniques = []
        
        # Разбиваем текст на предложения
        sentences = re.split(r'[.!?]+', text)
        
        for sentence in sentences:
            sentence = sentence.strip()
            if len(sentence) < 20 or len(sentence) > 300:
                continue
            
            # Проверяем каждую категорию техник
            for category, patterns in self.technique_patterns.items():
                for pattern in patterns:
                    if re.search(pattern, sentence, re.IGNORECASE):
                        quality = self.assess_quality(sentence)
                        
                        techniques.append({
                            'book': book_name,
                            'category': category,
                            'source': 'extracted',
                            'quality': quality,
                            'sales_technique': sentence
                        })
                        break
                else:
                    continue
                break
        
        return techniques
    
    def get_demo_techniques(self):
        """Добавляет базовые техники продаж"""
        demo_techniques = [
            {'book': 'theory_base', 'category': 'выявление_потребностей', 'source': 'theory', 'quality': 10, 'sales_technique': 'Что для вас важно при выборе квартиры?'},
            {'book': 'theory_base', 'category': 'выявление_потребностей', 'source': 'theory', 'quality': 10, 'sales_technique': 'Расскажите о ваших требованиях к жилью.'},
            {'book': 'theory_base', 'category': 'выявление_потребностей', 'source': 'theory', 'quality': 10, 'sales_technique': 'Какие факторы влияют на ваше решение?'},
            {'book': 'theory_base', 'category': 'выявление_потребностей', 'source': 'theory', 'quality': 10, 'sales_technique': 'Что должно быть в идеальной квартире?'},
            {'book': 'theory_base', 'category': 'выявление_потребностей', 'source': 'theory', 'quality': 10, 'sales_technique': 'Как вы представляете свой новый дом?'},
            {'book': 'theory_base', 'category': 'презентация_объектов', 'source': 'theory', 'quality': 10, 'sales_technique': 'Это означает, что вы сэкономите час времени каждый день.'},
            {'book': 'theory_base', 'category': 'презентация_объектов', 'source': 'theory', 'quality': 10, 'sales_technique': 'Благодаря этой планировке ваша семья будет чувствовать себя комфортно.'},
            {'book': 'theory_base', 'category': 'презентация_объектов', 'source': 'theory', 'quality': 10, 'sales_technique': 'Представьте, как здорово будет встречать рассветы на этом балконе.'},
            {'book': 'theory_base', 'category': 'презентация_объектов', 'source': 'theory', 'quality': 10, 'sales_technique': 'Эта особенность дает вам преимущество перед соседями.'},
            {'book': 'theory_base', 'category': 'работа_с_возражениями', 'source': 'theory', 'quality': 10, 'sales_technique': 'Понимаю ваши сомнения, давайте разберем детали.'},
            {'book': 'theory_base', 'category': 'работа_с_возражениями', 'source': 'theory', 'quality': 10, 'sales_technique': 'Согласен, цена важна, но посмотрите на выгоды.'},
            {'book': 'theory_base', 'category': 'работа_с_возражениями', 'source': 'theory', 'quality': 10, 'sales_technique': 'Многие клиенты сначала так думают, но потом понимают.'},
            {'book': 'theory_base', 'category': 'работа_с_возражениями', 'source': 'theory', 'quality': 10, 'sales_technique': 'Это важный вопрос, спасибо что подняли его.'},
            {'book': 'theory_base', 'category': 'установление_доверия', 'source': 'theory', 'quality': 10, 'sales_technique': 'Мой опыт показывает, что это лучшее решение.'},
            {'book': 'theory_base', 'category': 'установление_доверия', 'source': 'theory', 'quality': 10, 'sales_technique': 'Другие клиенты остались очень довольны.'},
            {'book': 'theory_base', 'category': 'установление_доверия', 'source': 'theory', 'quality': 10, 'sales_technique': 'Гарантирую, что вы не пожалеете о выборе.'},
            {'book': 'theory_base', 'category': 'установление_доверия', 'source': 'theory', 'quality': 10, 'sales_technique': 'Наша компания работает на рынке 15 лет.'},
            {'book': 'theory_base', 'category': 'закрытие_сделки', 'source': 'theory', 'quality': 10, 'sales_technique': 'Когда вам удобно подписать договор?'},
            {'book': 'theory_base', 'category': 'закрытие_сделки', 'source': 'theory', 'quality': 10, 'sales_technique': 'Готовы оформить бронирование сегодня?'},
            {'book': 'theory_base', 'category': 'закрытие_сделки', 'source': 'theory', 'quality': 10, 'sales_technique': 'Что выбираете - первый вариант или второй?'},
            {'book': 'theory_base', 'category': 'закрытие_сделки', 'source': 'theory', 'quality': 10, 'sales_technique': 'Когда планируете въезжать?'}
        ]
        return demo_techniques
    
    def assess_quality(self, text):
        """Оценивает качество техники продаж от 1 до 10"""
        quality = 5  # базовая оценка
        
        # Бонусы за полезные элементы
        if '?' in text:
            quality += 1  # вопросы важны
        if any(word in text.lower() for word in ['вы', 'ваш', 'вам']):
            quality += 1  # персонализация
        if len(text) > 50:
            quality += 1  # детализация
        if any(word in text.lower() for word in ['выгода', 'польза', 'экономия']):
            quality += 1  # фокус на выгодах
        
        # Штрафы
        if len(text) < 30:
            quality -= 1  # слишком короткие
        if len(text) > 200:
            quality -= 1  # слишком длинные
        
        return max(1, min(10, quality))
    
    def clean_text(self, text):
        """Очищает текст от лишних символов"""
        # Убираем HTML теги
        text = re.sub(r'<[^>]*>', ' ', text)
        # Убираем лишние пробелы
        text = re.sub(r'\s+', ' ', text)
        # Убираем спецсимволы
        text = re.sub(r'[^\w\s.,!?;:()\-«»""„"„'']', ' ', text)
        return text.strip()
    
    def save_to_csv(self, techniques):
        """Сохраняет техники в CSV файл"""
        csv_path = self.data_dir / "quality_techniques.csv"
        
        with open(csv_path, 'w', newline='', encoding='utf-8') as f:
            writer = csv.DictWriter(f, fieldnames=['book', 'category', 'source', 'quality', 'sales_technique'])
            writer.writeheader()
            writer.writerows(techniques)
        
        print(f"💾 Сохранено в: {csv_path}")
        return csv_path
    
    def save_summary(self, techniques):
        """Создает и сохраняет сводку по техникам"""
        summary = {
            'total_techniques': len(techniques),
            'categories': {},
            'books': {},
            'quality_stats': {
                'avg_quality': sum(t['quality'] for t in techniques) / len(techniques) if techniques else 0,
                'avg_length': sum(len(t['sales_technique']) for t in techniques) / len(techniques) if techniques else 0
            }
        }
        
        # Подсчет по категориям
        for technique in techniques:
            category = technique['category']
            book = technique['book']
            
            summary['categories'][category] = summary['categories'].get(category, 0) + 1
            summary['books'][book] = summary['books'].get(book, 0) + 1
        
        # Сохраняем сводку
        summary_path = self.data_dir / "quality_summary.json"
        with open(summary_path, 'w', encoding='utf-8') as f:
            json.dump(summary, f, ensure_ascii=False, indent=2)
        
        print(f"📋 Сводка сохранена в: {summary_path}")
        return summary


def main():
    """Главная функция обработки"""
    print("📚 ОБРАБОТКА ВСЕХ КНИГ ПО ПРОДАЖАМ")
    print("=" * 50)
    
    # Создаем процессор
    processor = SalesBookProcessor()
    
    # Обрабатываем все книги
    print("🔄 Запуск обработки книг...")
    techniques = processor.process_all_books()
    
    if techniques:
        # Сохраняем результаты
        processor.save_to_csv(techniques)
        summary = processor.save_summary(techniques)
        
        print("\n" + "="*50)
        print("📈 КРАТКАЯ СТАТИСТИКА:")
        print(f"Всего техник: {len(techniques)}")
        print(f"Средняя оценка качества: {summary['quality_stats']['avg_quality']:.1f}/10")
        print(f"Средняя длина: {summary['quality_stats']['avg_length']:.0f} символов")
        print("\nКатегории:")
        for category, count in summary['categories'].items():
            print(f"  {category}: {count}")
        
        print("\nКниги:")
        for book, count in summary['books'].items():
            if count > 0:
                print(f"  {book}: {count}")
    else:
        print("\n❌ Техники не найдены. Проверьте файлы в папке books/")


if __name__ == "__main__":
    main() 